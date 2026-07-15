# doc_tools.py -- LangChain @tool 定义：规范驱动 + 段落/run 原语 + 结构级操作
#
# 工具均为通用原语，不假定任何领域。apply_role_style 按激活的 StyleSpec 落样式，
# 其余为可直接调用的文档操作。execute_tools 与 /api/execute 共用 all_tools。
import os
import re
import tempfile
import threading
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from langchain_core.tools import tool

from spec_store import get_active_spec, ParaStyle, CompositeStyle

WORKSPACE_DOC = os.path.join(os.path.dirname(__file__), "workspace", "current.docx")
_doc_lock = threading.Lock()

# ---------- 内部辅助函数 ----------
def _save_doc(doc: Document) -> None:
    """原子写入：先写临时文件再 replace，避免并发/同路径保存损坏 zip。"""
    dest_dir = os.path.dirname(WORKSPACE_DOC)
    fd, tmp_path = tempfile.mkstemp(suffix=".docx", dir=dest_dir)
    os.close(fd)
    try:
        doc.save(tmp_path)
        os.replace(tmp_path, WORKSPACE_DOC)
    except Exception:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise

def _get_paragraph(idx: int, doc: Document):
    if idx < 0 or idx >= len(doc.paragraphs):
        raise IndexError(f"段落索引 {idx} 无效，文档共有 {len(doc.paragraphs)} 段")
    return doc.paragraphs[idx]

def _parse_indices(para_index) -> list:
    """兼容 int, list, 或 '1,2,3' 字符串格式的段落索引"""
    if isinstance(para_index, int):
        return [para_index]
    if isinstance(para_index, list):
        return [int(i) for i in para_index]
    if isinstance(para_index, str):
        return [int(i.strip()) for i in para_index.replace(" ", "").split(",") if i.strip()]
    raise ValueError(f"无效的 para_index 格式: {para_index}")

def _ensure_runs(para):
    """无 Run 的段落无法设置字体/加粗，补一个 Run。"""
    if not para.runs:
        para.add_run("")

# ---------- 字号解析 / CJK 排版辅助 ----------
_ALIGN_MAP = {
    "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
    "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
    "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
    "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
}

_CN_SIZE_RAW = {
    "初号": 42, "小初": 36, "一号": 26, "小一": 24, "二号": 22, "小二": 18,
    "三号": 16, "小三": 15, "四号": 14, "小四": 12, "五号": 10.5, "小五": 9,
    "六号": 7.5, "小六": 6.5, "七号": 5.5, "八号": 5,
}
_CN_DIGIT_MAP = str.maketrans("一二三四五六七八初", "123456780")

def _norm_size(text: str) -> str:
    text = text.strip()
    text = re.sub(r"(?i)pt$", "", text)
    text = text.replace(" ", "").replace("号", "")
    return text.translate(_CN_DIGIT_MAP)

CN_SIZE_TO_PT: dict[str, float] = {_norm_size(k): v for k, v in _CN_SIZE_RAW.items()}

def resolve_size(s) -> float | None:
    """接受 '小四'/'小5号'/'小 4 号'/'14pt'/'14'/14 → pt。"""
    if s is None:
        return None
    if isinstance(s, (int, float)):
        return float(s)
    text = str(s).strip()
    if not text:
        return None
    norm = _norm_size(text)
    if norm in CN_SIZE_TO_PT:
        return CN_SIZE_TO_PT[norm]
    try:
        return float(norm)
    except ValueError:
        raise ValueError(f"无法识别的字号: {s}")

def _para_font_size_pt(para) -> float:
    """读取段落实际字号：run > 段落样式 > 回退 10.5pt。"""
    for run in para.runs:
        if run.font.size is not None:
            return run.font.size.pt
    style = para.style
    if style is not None and style.font.size is not None:
        return style.font.size.pt
    return 10.5

def _ch_to_cm(ch: float, pt: float) -> float:
    """CJK 字符宽度换算：N 个字符 ≈ N × 字号pt；1pt = 2.54/72 cm。"""
    return ch * pt * 2.54 / 72

def _resolve_indent(s: str, pt: float):
    """缩进解析：'Nch' 按字号 pt 精确换算，'Ncm' 直用，纯数字按 cm。"""
    text = str(s).strip()
    if text.endswith("ch"):
        return Cm(_ch_to_cm(float(text.replace("ch", "")), pt))
    if text.endswith("cm"):
        return Cm(float(text.replace("cm", "")))
    return Cm(float(text))

def _resolve_color(color: str) -> RGBColor:
    if color.startswith("#"):
        r, g, b = bytes.fromhex(color[1:])
        return RGBColor(r, g, b)
    color_map = {
        "red": RGBColor(255, 0, 0), "blue": RGBColor(0, 0, 255),
        "green": RGBColor(0, 255, 0), "black": RGBColor(0, 0, 0),
    }
    return color_map.get(color.lower(), RGBColor(0, 0, 0))

def _set_rfonts(run, eastasia: str | None = None, ascii_: str | None = None):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    if eastasia:
        rFonts.set(qn("w:eastAsia"), eastasia)
    if ascii_:
        rFonts.set(qn("w:ascii"), ascii_)
        rFonts.set(qn("w:hAnsi"), ascii_)

# ---------- 样式应用（段落级 / run 级 / 复合切分） ----------
_RUN_ATTRS = ("font_east_asia", "font_ascii", "font_size", "bold", "color")

def _has_run_attrs(style: ParaStyle) -> bool:
    return any(getattr(style, a) is not None for a in _RUN_ATTRS)

def _apply_run_format(run, style: ParaStyle):
    if style.font_east_asia:
        _set_rfonts(run, eastasia=style.font_east_asia)
    if style.font_ascii:
        _set_rfonts(run, ascii_=style.font_ascii)
    if style.font_size is not None:
        run.font.size = Pt(style.font_size)
    if style.bold is not None:
        run.bold = style.bold
    if style.color is not None:
        run.font.color.rgb = _resolve_color(style.color)

def _apply_para_format(para, style: ParaStyle):
    pf = para.paragraph_format
    if style.alignment is not None:
        para.alignment = _ALIGN_MAP.get(style.alignment, para.alignment)
    if style.line_spacing is not None:
        pf.line_spacing = style.line_spacing
    if style.space_before is not None:
        pf.space_before = Pt(style.space_before)
    if style.space_after is not None:
        pf.space_after = Pt(style.space_after)
    eff_pt = style.font_size if style.font_size is not None else _para_font_size_pt(para)
    if style.first_line_indent is not None:
        pf.first_line_indent = _resolve_indent(style.first_line_indent, eff_pt)
    if style.left_indent is not None:
        pf.left_indent = _resolve_indent(style.left_indent, eff_pt)

def _apply_composite(para, comp: CompositeStyle):
    """按"引题/内容"切分段落 runs（确定性，不让 LLM 拆 run）。"""
    text = para.text
    m = re.match(comp.label_pattern, text)
    if not m:
        if _has_run_attrs(comp.content_style):
            _ensure_runs(para)
            for run in para.runs:
                _apply_run_format(run, comp.content_style)
        return
    label = m.group(0)
    rest = text[m.end():]
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    if label:
        _apply_run_format(para.add_run(label), comp.label_style)
    if rest:
        _apply_run_format(para.add_run(rest), comp.content_style)

def _coerce_style_dict(d: dict) -> dict:
    """把样式 dict 中的字号类字段统一解析为 pt（供 format_runs 使用）。"""
    out = dict(d)
    for k in ("font_size", "space_before", "space_after", "line_spacing"):
        if k in out and out[k] is not None:
            out[k] = resolve_size(out[k])
    return out

# ---------- 规范驱动工具 ----------
@tool
def apply_role_style(para_index: list[int], role: str) -> str:
    """按当前激活的排版规范，将 role 对应样式批量应用到段落（规范驱动，推荐优先使用）。
    role 为激活规范中定义的角色 id（见系统提示词角色列表）。相同角色的段落应合并为一次调用。
    规范未激活或角色未知时返回错误说明。"""
    spec = get_active_spec()
    if spec is None:
        return "未激活排版规范，无法使用 apply_role_style；请激活规范或改用 set_font 等原语工具。"
    rule = spec.roles.get(role)
    if rule is None:
        valid = ", ".join(spec.roles.keys())
        return f"未知角色: {role}。当前规范({spec.domain})可用角色: {valid}"
    with _doc_lock:
        doc = Document(WORKSPACE_DOC)
        indices = _parse_indices(para_index)
        for idx in indices:
            para = _get_paragraph(idx, doc)
            _apply_para_format(para, rule.style)
            if rule.composite is not None:
                _apply_composite(para, rule.composite)
            elif _has_run_attrs(rule.style):
                _ensure_runs(para)
                for run in para.runs:
                    _apply_run_format(run, rule.style)
        _save_doc(doc)
    return f"段落 {indices} 已应用角色 {role}({rule.label})，规范: {spec.domain}"

@tool
def format_runs(para_index: list[int], label_pattern: str,
                label_style: dict, content_style: dict) -> str:
    """按"引题/内容"切分段落 runs 分别格式化（用于规范未覆盖的复合段落手动处理）。
    label_pattern: 匹配引题前缀的正则；label_style/content_style: 样式字段，
    如 {'font_east_asia':'黑体','font_size':'小五'}。"""
    try:
        re.compile(label_pattern)
        comp = CompositeStyle(
            label_pattern=label_pattern,
            label_style=ParaStyle(**_coerce_style_dict(label_style)),
            content_style=ParaStyle(**_coerce_style_dict(content_style)),
        )
    except re.error as e:
        return f"正则无效: {e}"
    except Exception as e:
        return f"参数无效: {e}"
    with _doc_lock:
        doc = Document(WORKSPACE_DOC)
        indices = _parse_indices(para_index)
        for idx in indices:
            para = _get_paragraph(idx, doc)
            _apply_composite(para, comp)
        _save_doc(doc)
    return f"段落 {indices} 已按引题/内容切分格式化"

# ---------- 段落/run 级原语 ----------
@tool
def set_font(
    para_index: list[int],
    font_east_asia: Optional[str] = None,
    font_ascii: Optional[str] = None,
    font_size: Optional[str] = None,
) -> str:
    """批量修改段落的字体和字号。para_index 为段落索引数组，如 [0,1,2]。
    font_east_asia: 中文字体名，如"黑体"、"宋体"；font_ascii: 西文字体名；
    font_size: 字号名(如"小四")或 pt 数值/字符串(如"14"/"14pt")。"""
    with _doc_lock:
        doc = Document(WORKSPACE_DOC)
        indices = _parse_indices(para_index)
        size_pt = resolve_size(font_size) if font_size is not None else None
        for idx in indices:
            para = _get_paragraph(idx, doc)
            _ensure_runs(para)
            for run in para.runs:
                if font_east_asia:
                    _set_rfonts(run, eastasia=font_east_asia)
                if font_ascii:
                    _set_rfonts(run, ascii_=font_ascii)
                if size_pt is not None:
                    run.font.size = Pt(size_pt)
        _save_doc(doc)
    return f"段落 {indices} 字体已修改"

@tool
def set_bold(
    para_index: list[int],
    bold: bool = True,
) -> str:
    """批量设置段落加粗。para_index 为段落索引数组；bold 为 True 加粗，False 取消加粗。"""
    with _doc_lock:
        doc = Document(WORKSPACE_DOC)
        indices = _parse_indices(para_index)
        for idx in indices:
            para = _get_paragraph(idx, doc)
            _ensure_runs(para)
            for run in para.runs:
                run.bold = bold
        _save_doc(doc)
    return f"段落 {indices} 加粗: {bold}"

@tool
def set_alignment(
    para_index: list[int],
    alignment: str = "center",
) -> str:
    """批量设置段落对齐方式。para_index 为段落索引数组；alignment 取值: left, center, right, justify。"""
    if alignment not in _ALIGN_MAP:
        return f"不支持的对齐方式: {alignment}"
    with _doc_lock:
        doc = Document(WORKSPACE_DOC)
        indices = _parse_indices(para_index)
        for idx in indices:
            para = _get_paragraph(idx, doc)
            para.alignment = _ALIGN_MAP[alignment]
        _save_doc(doc)
    return f"段落 {indices} 对齐方式: {alignment}"

@tool
def set_indent(
    para_index: list[int],
    first_line_indent: Optional[str] = None,
    left_indent: Optional[str] = None,
    right_indent: Optional[str] = None,
) -> str:
    """批量设置段落缩进。para_index 为段落索引数组。
    first_line_indent: "2ch"(按段落字号精确换算)或"0.74cm"；left_indent/right_indent: 同样接受 ch 或 cm。"""
    with _doc_lock:
        doc = Document(WORKSPACE_DOC)
        indices = _parse_indices(para_index)
        for idx in indices:
            para = _get_paragraph(idx, doc)
            pf = para.paragraph_format
            pt = _para_font_size_pt(para)
            if first_line_indent:
                pf.first_line_indent = _resolve_indent(first_line_indent, pt)
            if left_indent:
                pf.left_indent = _resolve_indent(left_indent, pt)
            if right_indent:
                pf.right_indent = _resolve_indent(right_indent, pt)
        _save_doc(doc)
    return f"段落 {indices} 缩进已修改"

@tool
def set_spacing(
    para_index: list[int],
    line_spacing: Optional[str] = None,
    space_before: Optional[str] = None,
    space_after: Optional[str] = None,
) -> str:
    """批量设置段落间距。para_index 为段落索引数组。
    line_spacing: 行距(数字=倍数如1.5, 字符如"20pt"=固定值)；space_before/space_after: 段前/段后距，如"10pt"。"""
    with _doc_lock:
        doc = Document(WORKSPACE_DOC)
        indices = _parse_indices(para_index)
        for idx in indices:
            para = _get_paragraph(idx, doc)
            pf = para.paragraph_format
            if line_spacing is not None:
                if isinstance(line_spacing, str) and line_spacing.endswith("pt"):
                    pf.line_spacing = Pt(float(line_spacing.replace("pt", "")))
                else:
                    pf.line_spacing = float(line_spacing)
            if space_before:
                pf.space_before = Pt(float(str(space_before).replace("pt", "")))
            if space_after:
                pf.space_after = Pt(float(str(space_after).replace("pt", "")))
        _save_doc(doc)
    return f"段落 {indices} 间距已设置"

@tool
def set_color(
    para_index: list[int],
    color: str,
) -> str:
    """批量设置段落字体颜色。para_index 为段落索引数组；color 取值: #hex 如 #FF0000，或名称如 red/blue/green/black。"""
    rgb = _resolve_color(color)
    with _doc_lock:
        doc = Document(WORKSPACE_DOC)
        indices = _parse_indices(para_index)
        for idx in indices:
            para = _get_paragraph(idx, doc)
            _ensure_runs(para)
            for run in para.runs:
                run.font.color.rgb = rgb
        _save_doc(doc)
    return f"段落 {indices} 颜色已修改"

# ---------- 结构级工具 ----------
_PAGE_SIZES = {  # (宽 cm, 高 cm)
    "A4": (21.0, 29.7),
    "A5": (14.8, 21.0),
    "Letter": (21.59, 27.94),
    "Legal": (21.59, 35.56),
}

@tool
def set_page_setup(
    size: str = "A4",
    orientation: str = "portrait",
    margin_top: Optional[float] = None,
    margin_bottom: Optional[float] = None,
    margin_left: Optional[float] = None,
    margin_right: Optional[float] = None,
) -> str:
    """设置页面纸张大小、方向与页边距(单位 cm)。size: A4/A5/Letter/Legal；orientation: portrait/landscape。
    应用于全文所有节。"""
    if size not in _PAGE_SIZES:
        return f"不支持的纸张大小: {size}，可选: {list(_PAGE_SIZES.keys())}"
    w, h = _PAGE_SIZES[size]
    landscape = orientation.lower() == "landscape"
    with _doc_lock:
        doc = Document(WORKSPACE_DOC)
        for section in doc.sections:
            if landscape:
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width = Cm(h)
                section.page_height = Cm(w)
            else:
                section.orientation = WD_ORIENT.PORTRAIT
                section.page_width = Cm(w)
                section.page_height = Cm(h)
            if margin_top is not None:
                section.top_margin = Cm(margin_top)
            if margin_bottom is not None:
                section.bottom_margin = Cm(margin_bottom)
            if margin_left is not None:
                section.left_margin = Cm(margin_left)
            if margin_right is not None:
                section.right_margin = Cm(margin_right)
        _save_doc(doc)
    return f"页面已设置: {size}, {orientation}"

@tool
def apply_named_style(
    para_index: list[int],
    style_name: str,
    level: Optional[int] = None,
) -> str:
    """应用 Word 命名样式到段落（如 'Heading 1'/'Title'/'Normal'/'List Bullet'），
    并可设置大纲级别 level(0-9) 用于目录与导航窗格。"""
    with _doc_lock:
        doc = Document(WORKSPACE_DOC)
        try:
            style = doc.styles[style_name]
        except KeyError:
            names = [s.name for s in doc.styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
            return f"未找到样式: {style_name}。可用段落样式(部分): {names[:20]}"
        indices = _parse_indices(para_index)
        for idx in indices:
            para = _get_paragraph(idx, doc)
            para.style = style
            if level is not None:
                pPr = para._p.get_or_add_pPr()
                for old in pPr.findall(qn("w:outlineLvl")):
                    pPr.remove(old)
                ol = OxmlElement("w:outlineLvl")
                ol.set(qn("w:val"), str(int(level)))
                pPr.append(ol)
        _save_doc(doc)
    return f"段落 {indices} 已应用样式 {style_name}"

@tool
def insert_page_break(para_index: list[int]) -> str:
    """使指定段落另起一页开始（设置段前分页）。para_index 为段落索引数组。"""
    with _doc_lock:
        doc = Document(WORKSPACE_DOC)
        indices = _parse_indices(para_index)
        for idx in indices:
            para = _get_paragraph(idx, doc)
            para.paragraph_format.page_break_before = True
        _save_doc(doc)
    return f"段落 {indices} 已设为另起一页"

def _add_page_number_field(para):
    run = para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)

@tool
def set_header_footer(
    position: str,
    content_type: str,
    content: Optional[str] = None,
) -> str:
    """设置页眉或页脚。position: 'header'/'footer'；
    content_type: 'text'(写文字)/'page_number'(插入页码域)/'empty'(清空)；
    content: 当 content_type='text' 时的文字内容。应用于全文所有节。"""
    if position not in ("header", "footer"):
        return f"position 必须为 header 或 footer，收到: {position}"
    if content_type not in ("text", "page_number", "empty"):
        return f"content_type 必须为 text/page_number/empty，收到: {content_type}"
    with _doc_lock:
        doc = Document(WORKSPACE_DOC)
        for section in doc.sections:
            hf = section.header if position == "header" else section.footer
            hf.is_linked_to_previous = False
            para = hf.paragraphs[0] if hf.paragraphs else hf.add_paragraph()
            for r in list(para.runs):
                r._element.getparent().remove(r._element)
            if content_type == "text" and content:
                para.add_run(content)
            elif content_type == "page_number":
                _add_page_number_field(para)
        _save_doc(doc)
    return f"{position} 已设置: {content_type}"

# ---------- 工具列表 (供 execute_tools 与 /api/execute) ----------
all_tools = [
    apply_role_style, format_runs,
    set_font, set_bold, set_alignment, set_indent, set_spacing, set_color,
    set_page_setup, apply_named_style, insert_page_break, set_header_footer,
]
